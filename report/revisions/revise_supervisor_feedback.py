from __future__ import annotations

from copy import deepcopy
import re
import shutil
import struct
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SUBMISSION_SOURCE = ROOT / "report" / "Submission" / "1221303805_Arshad_Jamal_FYP1_Report.docx"
TRACKED_SOURCE = ROOT / "report" / "1221303805_Arshad_Jamal_FYP1_Report.docx"
SOURCE = SUBMISSION_SOURCE if SUBMISSION_SOURCE.exists() else TRACKED_SOURCE
OUT = ROOT / "report" / "Submission" / "1221303805_Arshad_Jamal_FYP1_Revised_Copy.docx"
APA_WORK = ROOT / "report" / "revisions" / "_apa_work"


def iter_block_items(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def find_paragraph(doc: Document, text: str, *, last: bool = False) -> Paragraph:
    matches = [p for p in doc.paragraphs if p.text.strip() == text]
    if not matches:
        raise ValueError(f"Paragraph not found: {text}")
    return matches[-1] if last else matches[0]


def find_startswith(doc: Document, prefix: str, *, last: bool = False) -> Paragraph:
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
    if not matches:
        raise ValueError(f"Paragraph starting with text not found: {prefix}")
    return matches[-1] if last else matches[0]


def first_text_paragraph_after(doc: Document, heading_text: str) -> Paragraph:
    blocks = list(iter_block_items(doc))
    start = next(
        i for i, block in enumerate(blocks)
        if isinstance(block, Paragraph) and block.text.strip() == heading_text
    )
    for block in blocks[start + 1 :]:
        if isinstance(block, Paragraph) and block.text.strip():
            return block
    raise ValueError(f"No paragraph after heading: {heading_text}")


def set_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def replace_inline_text(paragraph: Paragraph, old: str, new: str) -> None:
    replaced = False
    for node in paragraph._p.xpath(".//w:t"):
        if node.text and old in node.text:
            node.text = node.text.replace(old, new)
            replaced = True
    if not replaced and old in paragraph.text:
        set_text(paragraph, paragraph.text.replace(old, new))
        replaced = True
    if not replaced:
        raise ValueError(f"Text not found for inline replacement: {old}")


def delete_between(doc: Document, start_text: str, end_text: str) -> None:
    blocks = list(iter_block_items(doc))
    start = next(
        i for i, block in enumerate(blocks)
        if isinstance(block, Paragraph) and block.text.strip() == start_text
    )
    end = next(
        i for i, block in enumerate(blocks)
        if i > start and isinstance(block, Paragraph) and block.text.strip() == end_text
    )
    for block in blocks[start + 1 : end]:
        block._element.getparent().remove(block._element)


def add_paragraph_before(anchor: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    paragraph = anchor.insert_paragraph_before(text, style=style)
    return paragraph


def add_table_before(
    doc: Document,
    anchor: Paragraph,
    rows: list[list[str]],
    *,
    widths: list[float],
    header: bool = True,
    font_size: float = 9.5,
    first_column_labels: bool = False,
) -> Table:
    if not rows:
        raise ValueError("A table requires at least one row")
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    anchor._element.addprevious(table._tbl)
    format_table(
        table,
        widths=widths,
        header=header,
        font_size=font_size,
        first_column_labels=first_column_labels,
    )
    return table


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_table(
    table: Table,
    *,
    widths: list[float],
    header: bool,
    font_size: float,
    first_column_labels: bool = False,
) -> None:
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if col_index < len(widths):
                cell.width = Inches(widths[col_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            is_header = header and row_index == 0
            is_label = first_column_labels and col_index == 0
            if is_header:
                shade_cell(cell, "D9EAF7")
            elif is_label:
                shade_cell(cell, "EEF3F7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)
                    run.font.bold = bool(is_header or is_label)
    if header and table.rows:
        repeat_table_header(table.rows[0])


def set_table_data(
    table: Table,
    rows: list[list[str]],
    *,
    widths: list[float],
    font_size: float = 9.5,
) -> None:
    while table.rows:
        table._tbl.remove(table.rows[-1]._tr)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    format_table(table, widths=widths, header=True, font_size=font_size)


def table_by_headers(doc: Document, *headers: str) -> Table:
    for table in doc.tables:
        if not table.rows:
            continue
        row = [cell.text.strip() for cell in table.rows[0].cells]
        if row[: len(headers)] == list(headers):
            return table
    raise ValueError(f"Table not found with headers: {headers}")


def replace_abstract(doc: Document) -> None:
    delete_between(doc, "Abstract", "Table of Contents")
    anchor = find_paragraph(doc, "Table of Contents")
    anchor.paragraph_format.page_break_before = True
    paragraphs = [
        (
            "Cardiopulmonary recordings may contain heart sounds, lung sounds, sensor-contact noise, "
            "body movement and environmental artefacts in the same audio channel. Because the heart and "
            "lung components can overlap in time and frequency, fixed filtering alone may not produce "
            "usable separated outputs. This application-based project proposes a Machine Learning-Based "
            "System for Cardiopulmonary Sound Separation. Its purpose is to provide a reusable prototype "
            "workflow that accepts a mixed cardiopulmonary recording and produces separate heart and lung "
            "sound files. The system is limited to sound separation and does not perform disease diagnosis "
            "or provide clinical decisions."
        ),
        (
            "The FYP1 work combined a structured literature review, comparison of existing digital "
            "auscultation tools, requirements analysis, questionnaire data and software design. The "
            "literature review considered preprocessing, conventional decomposition methods, deep learning "
            "approaches, datasets and evaluation measures. A questionnaire collected 53 responses about "
            "processing challenges, preferred functions and interface expectations. Most respondents were "
            "from academic or technical backgrounds, so the results provide preliminary requirements rather "
            "than clinical validation. A targeted follow-up with healthcare staff and audio analysts is "
            "therefore required before the FYP2 requirements are finalised."
        ),
        (
            "The revised design treats the project as an application instead of a stand-alone algorithm. "
            "Healthcare staff can enter non-identifying audio attributes, upload recordings, follow the job "
            "status and access outputs. An audio analyst can review the audio file list, select a separation "
            "model, run processing, inspect history and review valid evaluation information. Chapter 4 "
            "defines the application architecture, use cases, data records and separation strategies, "
            "including fixed-filter, NMF, VMD and NeoSSNet-style approaches. Chapter 5 limits itself to the "
            "activities planned for FYP2: implementation, integration, testing, deployment and documentation. "
            "No separation accuracy, user-acceptance result or clinical performance is claimed at this stage."
        ),
    ]
    for text in paragraphs:
        add_paragraph_before(anchor, text)


def update_application_framing(doc: Document) -> None:
    set_text(
        first_text_paragraph_after(doc, "1.4 Project Scope"),
        "The current project scope is a prototype for cardiopulmonary sound separation. The application is designed around a practical clinic- or laboratory-oriented workflow: healthcare staff enter non-identifying recording attributes and upload one mixed cardiopulmonary audio file, while an audio analyst reviews the file record, selects a separation strategy, starts processing and inspects the generated outputs. The system validates and preprocesses the input before generating separate heart and lung sound files for preview and download.",
    )
    set_text(
        first_text_paragraph_after(doc, "1.7 Target Audience"),
        "The primary application roles are healthcare staff and audio analysts. Healthcare staff represent the person who prepares a recording, enters non-identifying audio attributes and uploads the file. The audio analyst represents the person who reviews the audio list, selects a separation model, runs processing and examines the separated outputs. Students and researchers remain secondary users for development and evaluation, but the workflow is no longer described as being for any unspecified user.",
    )


def add_chapter2_application_study(doc: Document) -> None:
    set_text(
        first_text_paragraph_after(doc, "2.7 Existing Systems and Prototype Relevance"),
        "Existing work includes research platforms and commercial digital-auscultation applications. Their main value to this FYP is not that they already separate mixed heart and lung sounds, but that they show how body-sound recordings are captured, described, organised, reviewed and shared in a usable application. Eko supports recording, playback, annotation, waveform or PCG viewing, storage and sharing functions, while the Littmann CORE device uses the Eko application for visualisation and recording support (3M, n.d.; Eko Health, 2025). Thinklabs Wave provides real-time waveform and frequency displays together with recording, annotation, file browsing and sharing features (Thinklabs Medical LLC, n.d.). StethAid extends this idea into a platform with a digital stethoscope, mobile applications, web portals and centralised storage (Arjoune et al., 2023).",
    )
    second = find_startswith(doc, "This finding is critical to the current research")
    set_text(
        second,
        "These applications demonstrate that a useful solution needs more than a model execution button. Recording attributes, a clear audio-file list, traceable processing status, result review and controlled output handling are part of the application workflow. The proposed FYP does not copy diagnostic or patient-management functions from these products. Instead, it adopts the relevant software features and adds the project-specific contribution: selectable heart-lung sound separation strategies that produce separate audio outputs.",
    )

    anchor = find_paragraph(doc, "2.8 Literature Review Matrix")
    add_paragraph_before(anchor, "2.7.1 Comparative Study of Existing Tools and Applications", "Heading 3")
    add_paragraph_before(
        anchor,
        "The following comparison identifies the application features that are relevant to the proposed prototype and the gap that remains for this project.",
    )
    caption = add_paragraph_before(
        anchor,
        "Table 901: Comparative study of cardiopulmonary sound tools and applications",
        "Caption",
    )
    caption.paragraph_format.keep_with_next = True
    add_table_before(
        doc,
        anchor,
        [
            ["Tool or application", "Main workflow", "Observed features", "Relevance and limitation for this FYP"],
            [
                "Eko App",
                "Digital auscultation and recording review",
                "Record, play back, annotate and save body sounds; display PCG waveforms; store and share recordings (Eko Health, 2025).",
                "Supports preview, notes, visualisation and record management. It does not provide the proposed selectable heart-lung separation workflow.",
            ],
            [
                "3M Littmann CORE with Eko App",
                "Device-assisted digital auscultation",
                "Visualise, record, save, annotate and share sounds through the connected application (3M, n.d.).",
                "Shows a practical healthcare recording workflow. Device integration and diagnosis functions remain outside this FYP scope.",
            ],
            [
                "Thinklabs Wave",
                "Capture and review digital body sounds",
                "Waveform and frequency display, recording, annotation, file browsing, storage and sharing (Thinklabs Medical LLC, n.d.).",
                "Supports the need for audio lists, attributes, preview and optional waveform or spectrogram displays.",
            ],
            [
                "StethAid",
                "Pediatric digital-auscultation platform",
                "Digital stethoscope, mobile applications, user portals, sound libraries and centralised storage (Arjoune et al., 2023).",
                "Shows how roles, records and storage form an application. Its clinical analysis functions are not adopted by this separation-only prototype.",
            ],
            [
                "Proposed FYP solution",
                "Role-based upload and heart-lung sound separation",
                "Audio attributes, file validation, audio list, model selection, processing status, separate outputs, preview, download, history and valid metrics.",
                "Adds an end-to-end, model-selectable separation workflow while avoiding diagnosis claims.",
            ],
        ],
        widths=[1.25, 1.35, 1.9, 2.0],
        font_size=8.5,
    )

    add_paragraph_before(anchor, "2.7.2 Feature Decisions for the Proposed Solution", "Heading 3")
    add_paragraph_before(
        anchor,
        "The comparison was translated into a feature decision so that the proposed system remains achievable as an FYP prototype. Core features support the complete application flow. Additional features can be developed after that flow is stable, while diagnosis and treatment functions are excluded.",
    )
    caption = add_paragraph_before(
        anchor,
        "Table 902: Feature decisions for the proposed cardiopulmonary sound separation solution",
        "Caption",
    )
    caption.paragraph_format.keep_with_next = True
    add_table_before(
        doc,
        anchor,
        [
            ["Feature", "Decision", "Reason"],
            ["Audio attributes and upload", "Core FYP2", "Provides context for each recording and supports a healthcare-staff workflow."],
            ["File validation and correction messages", "Core FYP2", "Prevents unsupported or unreadable audio from entering the separation process."],
            ["Audio file list and filters", "Core FYP2", "Allows an audio analyst to find records and review their processing state."],
            ["Model selection and separation", "Core FYP2", "Implements the main project contribution and supports strategy comparison."],
            ["Status, preview, download and history", "Core FYP2", "Completes the user workflow and makes outputs traceable."],
            ["Waveform or spectrogram display", "Additional FYP2", "Improves review but is secondary to reliable output generation."],
            ["Notes, annotations and audit information", "Additional FYP2", "Supports handover between roles after the core workflow is stable."],
            ["Evaluation metrics", "Conditional", "Display only when reference data or a valid evaluation procedure is available."],
            ["Disease diagnosis or treatment advice", "Excluded", "Outside the approved sound-separation scope and not supported by this prototype."],
        ],
        widths=[2.0, 1.2, 3.3],
        font_size=9.0,
    )

    set_text(
        first_text_paragraph_after(doc, "2.10 Summary"),
        "This review connected cardiopulmonary sound research with the design of an application-based solution. The technical studies support noise-aware preprocessing, conventional baselines, machine learning strategies, public datasets and objective evaluation. The comparison of Eko, Littmann CORE, Thinklabs Wave and StethAid added an application perspective: useful systems organise recordings, preserve attributes, provide file lists, show progress, support review and keep outputs traceable. These findings directly inform the role-based requirements in Chapter 3 and the system design in Chapter 4.",
    )


def update_requirements_chapter(doc: Document) -> None:
    set_text(
        first_text_paragraph_after(doc, "3.1 Overview"),
        "The requirements for the Machine Learning-Based System for Cardiopulmonary Sound Separation were identified from the project scope, literature and document review, comparison of existing applications, supervisor consultation and a questionnaire. The Google Form received 53 responses and provided useful early information about processing challenges, preferred features and interface expectations. However, most respondents came from academic or technical backgrounds rather than hospital or clinic roles. The results are therefore treated as preliminary requirements, and a targeted follow-up with healthcare staff and audio analysts is required before the FYP2 requirements are finalised.",
    )
    set_text(
        first_text_paragraph_after(doc, "3.2 Questionnaire"),
        "The questionnaire was used to collect practical preferences for the prototype. It covered respondent background, cardiopulmonary processing challenges, system features, interface style, output presentation and evaluation needs. Ten requirement-focused items are analysed in this chapter, while the complete 24-question instrument remains in Appendix H.",
    )
    design = first_text_paragraph_after(doc, "3.2.1 Questionnaire Design")
    set_text(
        design,
        "The questionnaire was distributed through Google Forms and received 53 responses. The initial target included Software Engineering, Computer Science, artificial intelligence, machine learning, biomedical engineering and health-related students, researchers and supervisors. This convenience sample was suitable for early prototype feedback, but it did not adequately represent the healthcare staff and audio analysts expected to use a clinic- or laboratory-oriented application.",
    )
    next_design = find_startswith(doc, "The questionnaire was posted in Google Forms")
    set_text(
        next_design,
        "The introduction explained the academic purpose of the study, confidentiality and the absence of right or wrong answers. Multiple-choice, checkbox, Likert-scale and paragraph questions were used. The collected responses support preliminary feature and usability decisions. They do not constitute clinical validation or proof that the prototype is ready for hospital use.",
    )
    profile = find_startswith(doc, "With respect to the level of expertise")
    set_text(
        profile,
        "With respect to familiarity with audio or signal processing, 29 respondents (54.7%) selected basic, 15 (28.3%) selected not familiar, 6 (11.3%) selected intermediate and 3 (5.7%) selected advanced. Thirty respondents (56.6%) reported prior biomedical-audio experience, while 22 (41.5%) reported no prior experience. The sample therefore provides useful feedback from technical and student users, including some health-related respondents, but it cannot be treated as representative of hospital personnel. That limitation affects how confidently the requirements can be generalised to a clinic workflow.",
    )

    anchor = find_paragraph(doc, "3.3 Functional Requirements")
    add_paragraph_before(anchor, "3.2.13 Target-User Validation Limitation and Follow-Up", "Heading 3")
    add_paragraph_before(
        anchor,
        "The 53-response questionnaire remains useful for identifying general expectations such as simple upload, clear error messages, output preview, download and evaluation information. It does not fully answer role-specific questions about who records the audio, which attributes must be entered, how files are handed to an analyst, or how previous recordings are retrieved in a healthcare setting.",
    )
    add_paragraph_before(
        anchor,
        "Before the FYP2 requirements are frozen, a purposive follow-up should be conducted with healthcare personnel who handle or review body-sound recordings and with audio analysts or biomedical-signal researchers. The follow-up should confirm the minimum audio attributes, the audio-list workflow, role responsibilities, privacy expectations, result presentation and acceptable error messages. No responses from these target roles are invented in this report.",
    )

    functional = table_by_headers(doc, "Req ID", "Requirement", "Description")
    set_table_data(
        functional,
        [
            ["Req ID", "Requirement", "Description"],
            ["F1", "Enter Audio Attributes", "The system shall allow healthcare staff to enter non-identifying attributes such as recording source, body location, device, sample rate, duration and notes."],
            ["F2", "Upload Audio File", "The system shall allow healthcare staff or an audio analyst to upload a mixed cardiopulmonary WAV file."],
            ["F3", "Validate Audio and Attributes", "The system shall validate the file type, header, readability, size and required attributes before processing."],
            ["F4", "List Audio Files", "The system shall display audio records with filename, attributes, upload time, selected model, status and result availability."],
            ["F5", "Search or Filter Audio Records", "The system shall allow the audio analyst to filter records by filename, status, model or supported attributes."],
            ["F6", "Preprocess Audio", "The system shall apply the preprocessing required by the selected separation strategy."],
            ["F7", "Select Separation Model", "The system shall allow the audio analyst to select an available separation model or use the active default."],
            ["F8", "Run Sound Separation", "The system shall execute the selected strategy on the validated and preprocessed recording."],
            ["F9", "Generate Heart Sound Output", "The system shall create a separated heart sound file."],
            ["F10", "Generate Lung Sound Output", "The system shall create a separated lung sound file."],
            ["F11", "Display Processing Status", "The system shall show validation, preprocessing, running, completed or failed status with a clear message."],
            ["F12", "Preview Separated Outputs", "The system shall allow authorised users to preview the heart and lung outputs."],
            ["F13", "Download Separated Outputs", "The system shall allow authorised users to download the generated heart and lung files."],
            ["F14", "View Processing History", "The system shall retain a traceable list of audio records, jobs, selected models and output locations."],
            ["F15", "Display Valid Evaluation Information", "The system shall display evaluation metrics only when reference data or a valid evaluation procedure is available."],
        ],
        widths=[0.65, 1.65, 4.2],
        font_size=9.0,
    )

    non_functional = table_by_headers(doc, "Req ID", "Category", "Requirement")
    set_table_data(
        non_functional,
        [
            ["Req ID", "Category", "Requirement"],
            ["NF1", "Usability", "Healthcare staff shall be able to enter attributes and upload a file without configuring algorithm details."],
            ["NF2", "Performance", "The prototype shall show progress and avoid appearing unresponsive during processing."],
            ["NF3", "Reliability", "A failed job shall not be reported as completed, and incomplete outputs shall not be offered for download."],
            ["NF4", "Maintainability", "Validation, preprocessing, model execution, storage and result handling shall remain separate modules."],
            ["NF5", "Security and Privacy", "The prototype shall avoid unnecessary identifying information and restrict file access to the local application workflow."],
            ["NF6", "Compatibility", "The first prototype shall support WAV input and produce common downloadable audio output files."],
            ["NF7", "Data Integrity", "Audio attributes, selected model, job status and output paths shall remain linked to the correct record."],
            ["NF8", "Scalability", "The design shall allow more models and records to be added without changing the complete workflow."],
        ],
        widths=[0.65, 1.45, 4.4],
        font_size=9.0,
    )

    user_requirements = table_by_headers(doc, "User Requirement ID", "User Need", "Description")
    set_table_data(
        user_requirements,
        [
            ["User Requirement ID", "User Need", "Description"],
            ["UR1", "Healthcare staff enter recording details", "Healthcare staff need to record non-identifying audio attributes before upload."],
            ["UR2", "Healthcare staff upload a recording", "Healthcare staff need a simple upload process with clear validity feedback."],
            ["UR3", "Audio analyst reviews the audio list", "The analyst needs a searchable list of records, attributes and processing status."],
            ["UR4", "Audio analyst selects a model", "The analyst needs to choose an available strategy or use the active default model."],
            ["UR5", "Audio analyst starts separation", "The analyst needs to run processing for a selected audio record and model."],
            ["UR6", "Users understand job status", "Both roles need clear progress, completion and error messages."],
            ["UR7", "Users preview and download outputs", "Both roles need to access separate heart and lung files after successful processing."],
            ["UR8", "Audio analyst reviews history and metrics", "The analyst needs traceable previous jobs and valid evaluation information when available."],
        ],
        widths=[0.9, 1.9, 3.7],
        font_size=9.0,
    )

    set_text(
        first_text_paragraph_after(doc, "3.6 Summary"),
        "This chapter used the 53 questionnaire responses as preliminary input and combined them with the literature, existing-application comparison, project scope and supervisor feedback. The requirements now define two concrete roles: healthcare staff and audio analyst. They also add audio attributes, an audio file list, filtering, role-specific model selection, traceable history and conditional evaluation information. Because the original questionnaire did not adequately represent hospital or clinic personnel, targeted validation with the intended users remains an FYP2 requirement rather than a completed result.",
    )


USE_CASES = [
    {
        "id": "UC01",
        "name": "Enter audio file attributes",
        "actor": "Healthcare Staff",
        "pre": "The upload form is available and no patient-identifying information is required.",
        "trigger": "Healthcare staff prepares a recording for upload.",
        "flow": "1. Open the upload form. 2. Enter supported attributes such as body location, recording device and notes. 3. Continue to audio selection.",
        "alternative": "Missing required attributes are highlighted before submission.",
        "post": "The attributes are ready to be saved with the uploaded audio record.",
    },
    {
        "id": "UC02",
        "name": "Upload mixed cardiopulmonary audio",
        "actor": "Healthcare Staff",
        "pre": "A supported WAV file and the required non-identifying attributes are available.",
        "trigger": "Healthcare staff submits the upload form.",
        "flow": "1. Select the WAV file. 2. Submit the form. 3. The system validates the file. 4. The system stores the file and creates an audio record.",
        "alternative": "An unsupported, unreadable or oversized file is rejected with a correction message.",
        "post": "A valid audio record appears in the audio file list with an uploaded status.",
    },
    {
        "id": "UC03",
        "name": "View audio file list",
        "actor": "Audio Analyst; Healthcare Staff (limited view)",
        "pre": "At least one audio record has been uploaded.",
        "trigger": "The user opens the audio records page.",
        "flow": "1. Request the record list. 2. The system returns filenames, attributes, status and result availability. 3. The user filters or selects a record.",
        "alternative": "If no records exist, the system displays an empty-state message rather than an error.",
        "post": "A selected audio record is available for processing or review.",
    },
    {
        "id": "UC04",
        "name": "Select separation model",
        "actor": "Audio Analyst",
        "pre": "The selected audio record is valid and at least one model definition is active.",
        "trigger": "The analyst opens the model selector for an audio record.",
        "flow": "1. Load active model definitions. 2. Display model name, version and strategy type. 3. Store the analyst's selection with the job request.",
        "alternative": "If no model is selected, the active default is used. If no active model exists, processing is blocked with an error.",
        "post": "A valid model is associated with the pending separation job.",
    },
    {
        "id": "UC05",
        "name": "Run sound separation",
        "actor": "Audio Analyst",
        "pre": "A valid audio record and model selection are available.",
        "trigger": "The analyst selects Run Separation.",
        "flow": "1. Create a job. 2. Preprocess the audio. 3. Resolve the configured strategy. 4. Run separation. 5. Save heart and lung outputs. 6. Mark the job completed.",
        "alternative": "Validation, preprocessing, model or output failure marks the job failed and records an explanatory message.",
        "post": "The job is completed with two output paths, or failed with traceable error information.",
    },
    {
        "id": "UC06",
        "name": "View processing status",
        "actor": "Healthcare Staff and Audio Analyst",
        "pre": "A processing job exists.",
        "trigger": "The user opens or refreshes the job page.",
        "flow": "1. Request the job state. 2. Display validation, preprocessing, running, completed or failed status. 3. Show the associated message and selected model.",
        "alternative": "If the job cannot be found, the system displays a record-not-found message.",
        "post": "The user understands the current job state and available next action.",
    },
    {
        "id": "UC07",
        "name": "Preview separated outputs",
        "actor": "Healthcare Staff and Audio Analyst",
        "pre": "The job is completed and both output files exist.",
        "trigger": "The user opens the completed result page.",
        "flow": "1. Load result metadata. 2. Display separate heart and lung audio players. 3. Allow the user to listen to each output.",
        "alternative": "A missing output disables its player and reports that the result is incomplete.",
        "post": "The user has reviewed the available outputs before download.",
    },
    {
        "id": "UC08",
        "name": "Download separated outputs",
        "actor": "Healthcare Staff and Audio Analyst",
        "pre": "The requested output exists for a completed job.",
        "trigger": "The user selects Download Heart or Download Lung.",
        "flow": "1. Verify the job and output type. 2. Resolve the stored file path. 3. Return the selected audio file.",
        "alternative": "An unavailable output returns a clear message and no empty file is downloaded.",
        "post": "The selected heart or lung sound file is available on the user's device.",
    },
    {
        "id": "UC09",
        "name": "View processing history",
        "actor": "Audio Analyst",
        "pre": "Previous audio records or jobs exist.",
        "trigger": "The analyst opens the processing-history view.",
        "flow": "1. Request previous jobs. 2. Display audio attributes, model, status and completion time. 3. Open a selected result or failure record.",
        "alternative": "Filters that return no matches display an empty result set.",
        "post": "The analyst can trace previous processing activity without repeating the job.",
    },
    {
        "id": "UC10",
        "name": "View evaluation metrics",
        "actor": "Audio Analyst",
        "pre": "The job is completed and valid reference data or an approved metric calculation is available.",
        "trigger": "The analyst opens the evaluation area of a result.",
        "flow": "1. Retrieve stored metrics. 2. Display metric name, value and context. 3. Keep the audio outputs available for comparison.",
        "alternative": "If valid metrics are unavailable, the system states that no evaluation result is available and does not fabricate a score.",
        "post": "The analyst can review valid technical evaluation information without a diagnosis claim.",
    },
]


def rebuild_use_case_descriptions(doc: Document) -> None:
    delete_between(doc, "4.3.1 Use Case Description", "4.4 Activity Diagram")
    anchor = find_paragraph(doc, "4.4 Activity Diagram")
    add_paragraph_before(anchor, "4.3.1.1 Actors", "Heading 4")
    add_paragraph_before(
        anchor,
        "Healthcare Staff is responsible for entering non-identifying recording attributes, uploading audio, following status and accessing completed outputs. Audio Analyst is responsible for reviewing the audio list, selecting a model, starting separation, checking technical results and reviewing history. The roles describe application responsibilities; they do not imply that the prototype performs diagnosis.",
    )
    add_paragraph_before(anchor, "4.3.1.2 Use Case Description Tables", "Heading 4")
    add_paragraph_before(
        anchor,
        "Each use case in the diagram is described separately below so that its actor, conditions, main flow and exception handling are explicit.",
    )
    for index, item in enumerate(USE_CASES, start=1):
        caption = add_paragraph_before(
            anchor,
            f"Table {910 + index}: Use case description for {item['id']} - {item['name']}",
            "Caption",
        )
        caption.paragraph_format.keep_with_next = True
        add_table_before(
            doc,
            anchor,
            [
                ["Use case ID", item["id"]],
                ["Use case name", item["name"]],
                ["Primary actor", item["actor"]],
                ["Preconditions", item["pre"]],
                ["Trigger", item["trigger"]],
                ["Main flow", item["flow"]],
                ["Alternative / exception", item["alternative"]],
                ["Postconditions", item["post"]],
            ],
            widths=[1.55, 4.95],
            header=False,
            first_column_labels=True,
            font_size=9.0,
        )


def move_algorithm_design_to_chapter4(doc: Document) -> None:
    blocks = list(iter_block_items(doc))
    start = next(
        i for i, block in enumerate(blocks)
        if isinstance(block, Paragraph) and block.text.strip() == "5.1.3 Machine Learning Model Integration"
    )
    end = next(
        i for i, block in enumerate(blocks)
        if i > start and isinstance(block, Paragraph) and block.text.strip() == "5.1.4 Database and Storage Integration"
    )
    source_blocks = blocks[start + 1 : end]
    target = find_paragraph(doc, "4.5.3 Database and Storage Design")
    add_paragraph_before(target, "4.5.3 Separation Algorithm Design and Mathematical Formulation", "Heading 3")
    add_paragraph_before(
        target,
        "The algorithm design is placed in Chapter 4 because it defines how the proposed solution will transform one mixed recording into estimated heart and lung outputs. The strategies share one application workflow and one output contract, but they differ in how they represent and separate the signal. The equations and strategy descriptions below are design specifications for FYP2, not completed performance results.",
    )

    heading_map = {
        "5.1.3.1 Input Preparation": "4.5.3.1 Input Preparation",
        "5.1.3.2 Algorithm 1: Fixed Filter Baseline": "4.5.3.2 Fixed Filter Baseline",
        "5.1.3.3 Algorithm 2: NMF and VMD-Based Decomposition": "4.5.3.3 NMF and VMD-Based Decomposition",
        "5.1.3.4 Algorithm 3: NeoSSNet Deep Separation Strategy": "4.5.3.4 NeoSSNet Deep Separation Strategy",
        "5.1.3.5 Model Selection Strategy": "4.5.3.5 Model Selection Strategy",
    }
    skipped_intro = False
    for block in source_blocks:
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not skipped_intro and text.startswith("The machine learning model will be integrated"):
                skipped_intro = True
                continue
            if text in heading_map:
                add_paragraph_before(target, heading_map[text], block.style.name)
                continue
        target._element.addprevious(deepcopy(block._element))

    set_text(target, "4.5.4 Database and Storage Design")
    delete_between(doc, "5.1.3 Machine Learning Model Integration", "5.1.4 Database and Storage Integration")
    fyp2_anchor = find_paragraph(doc, "5.1.4 Database and Storage Integration")
    add_paragraph_before(
        fyp2_anchor,
        "The separation algorithms and their mathematical design are defined in Chapter 4. The FYP2 activity is to implement the shared strategy interface, connect the selected baseline and machine learning strategies, and verify that each configured strategy returns separate heart and lung output files through the same application workflow.",
    )
    add_paragraph_before(
        fyp2_anchor,
        "Implementation will begin with an end-to-end path using a baseline strategy. NMF, VMD and a NeoSSNet-style strategy can then be integrated according to data, hardware and supervisor approval. Training, fine-tuning and final metric reporting will be completed only when valid datasets and reference signals are available.",
    )


def update_design_and_plan(doc: Document) -> None:
    set_text(
        first_text_paragraph_after(doc, "4.1 Overview"),
        "This chapter defines an application-oriented design for the cardiopulmonary sound separation prototype. The design separates responsibilities between Healthcare Staff and Audio Analyst and covers audio attributes, upload, file listing, model selection, processing, output review, download and history. The diagrams and use cases describe a planned FYP2 solution rather than a completed hospital product, and no diagnostic function is included.",
    )
    set_text(
        first_text_paragraph_after(doc, "4.2 Context Diagram"),
        "The context diagram shows the proposed system and its external interactions. Healthcare staff enter non-identifying audio attributes, upload a mixed recording and access status or outputs. The audio analyst reviews audio records, selects a model, runs separation and inspects technical results. Approved public cardiopulmonary datasets support development and evaluation. The system boundary remains limited to heart-lung sound separation.",
    )
    set_text(
        first_text_paragraph_after(doc, "4.3 Use Case Diagram"),
        "The use case diagram replaces the previous unspecified actor with two concrete roles. Healthcare Staff handles recording entry and upload, while Audio Analyst handles model selection, separation and technical review. Shared use cases cover status, preview and download. The diagram uses ten use cases that are described individually in the following tables.",
    )
    rebuild_use_case_descriptions(doc)
    move_algorithm_design_to_chapter4(doc)

    set_text(
        first_text_paragraph_after(doc, "4.4 Activity Diagram"),
        "The activity diagram follows the role-based workflow from audio preparation to output access. Healthcare staff enter attributes and upload a recording. The system validates and stores the audio record. The audio analyst then selects the record and model, after which the system preprocesses the file and runs the chosen strategy. Completion requires both heart and lung output files; otherwise, the job is marked as failed with an explanatory message.",
    )
    set_text(
        first_text_paragraph_after(doc, "4.6 Sequence Diagram"),
        "The sequence diagram focuses on the Audio Analyst's model-selection and separation use case. Every lifeline corresponds to a class or interface defined in the class diagram. PrototypeInterface passes the selected audio record and model to AnalysisController. ModelRegistry supplies the model definition, SeparationService creates and runs the job, StrategyFactory resolves a SeparationStrategy, and ResultService stores the completed result for preview, download or later review.",
    )
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Caption" and "Sequence diagram for uploading and separating" in paragraph.text:
            set_text(
                paragraph,
                re.sub(
                    r"Sequence diagram for uploading and separating a cardiopulmonary recording",
                    "Sequence diagram for model selection and cardiopulmonary sound separation",
                    paragraph.text,
                ),
            )
    interface = first_text_paragraph_after(doc, "4.7 Interface Design")
    set_text(
        interface,
        "The interface is organised around the two application roles. The upload view provides non-identifying audio attribute fields, a WAV selector and validation feedback for healthcare staff. The audio analyst view provides an audio file list, filters, model selection, processing status, result preview, download controls, history and valid evaluation information. The wording must consistently describe sound separation rather than diagnosis.",
    )
    set_text(
        first_text_paragraph_after(doc, "4.8 Summary"),
        "The revised system design treats the project as a role-based application. It specifies how healthcare staff and audio analysts work with audio records, how each use case behaves, how data and results remain traceable, and how multiple separation strategies fit behind one interface. The detailed algorithm design is now contained in this chapter. Chapter 5 therefore focuses only on the implementation, testing, deployment and documentation activities planned for FYP2.",
    )

    set_text(
        first_text_paragraph_after(doc, "5.1 Development Phase"),
        "FYP2 will implement the application design defined in Chapter 4. Work will begin with the role-based audio-record workflow, including attributes, upload, validation and the audio file list. The backend, storage and model registry will then be connected to a separation strategy. Result preview, download, history and conditional evaluation information will be added after the core path is stable. This chapter describes planned activities only and does not claim completed model or user-test results.",
    )
    set_text(
        first_text_paragraph_after(doc, "5.1.1 Front-End Development"),
        "The front end will provide separate task paths for healthcare staff and audio analysts within one prototype interface. Healthcare staff need a simple audio attribute and upload form. Audio analysts need a record list, filters, model selection, job controls and result review. Shared views will provide status, preview and download functions. The interface will avoid diagnosis wording because the prototype separates sounds only.",
    )
    replacements = {
        "The upload area: a file selector for WAV file upload and process initiation": "The upload area: non-identifying audio attribute fields, a WAV selector and validation feedback",
        "The model selector: an element allowing the user to switch between separation algorithms or default models": "The audio file list: searchable records with attributes, status, selected model and result availability",
        "The process status area: information about the upload progress, validation status, model initiation, process progress, and possible errors": "The analyst controls: model selection, process initiation, status and clear failure messages",
        "The result preview area: a preview of separated heart and lung sounds and respective download buttons": "The result area: separate heart and lung previews, download buttons and valid metric information",
        "The job history area: a list of completed processes with the possibility to re-open them": "The history area: previous jobs that can be filtered and reopened for review",
    }
    for old, new in replacements.items():
        try:
            set_text(find_paragraph(doc, old), new)
        except ValueError:
            pass

    uat = first_text_paragraph_after(doc, "5.2.4 User Acceptance Testing")
    set_text(
        uat,
        "User acceptance testing will be conducted after the prototype workflow is stable. Healthcare staff will be asked to complete audio-attribute entry, upload, status and output-access tasks. Audio analysts will be asked to find a record, select a model, run separation, review results and reopen a previous job. These activities are planned for FYP2 and have not yet been completed.",
    )
    set_text(
        first_text_paragraph_after(doc, "5.5 Summary"),
        "Chapter 5 now describes the FYP2 activities at implementation-plan level. The work covers the role-based front end, backend services, audio attributes, file listing, database and storage integration, strategy integration, planned testing, deployment and documentation. Technical algorithm details have been moved to Chapter 4. No separation quality, testing outcome or healthcare acceptance result is reported as completed.",
    )

    set_text(
        first_text_paragraph_after(doc, "6.2 Project Summary"),
        "The FYP1 report establishes the basis for a role-based cardiopulmonary sound separation application. It connects the literature and questionnaire findings to a workflow in which healthcare staff prepare and upload audio records while an audio analyst selects a model, runs separation and reviews traceable outputs. The design remains a prototype specification for FYP2 rather than a completed clinical system.",
    )
    set_text(
        first_text_paragraph_after(doc, "6.3 Expected Contributions"),
        "The expected contribution is a reusable application workflow for managing and separating mixed cardiopulmonary recordings. In addition to producing heart and lung output files, the prototype is expected to manage audio attributes, validation, file lists, model selection, status, history and result access. The strategy interface also allows fixed-filter, NMF, VMD and NeoSSNet-style approaches to be compared without redesigning the surrounding application.",
    )
    set_text(
        first_text_paragraph_after(doc, "6.4 Limitations"),
        "The current work is limited to FYP1 design and planning. The 53 questionnaire responses mainly represent academic and technical respondents, so they do not validate a hospital workflow. Model implementation, separation performance, healthcare-user validation, security testing and deployment remain incomplete. The prototype is also limited to sound separation and must not be presented as a diagnostic system.",
    )
    set_text(
        first_text_paragraph_after(doc, "6.5 Future Work"),
        "FYP2 should first confirm the revised requirements with healthcare staff and audio analysts. Development can then implement audio attributes, the audio file list, model selection and the complete output workflow before improving the separation strategies. Later work may add waveform or spectrogram displays, annotations and stronger access control after the core application is working and has been tested honestly.",
    )

    # The final issue metadata assigns the phase-enhanced transformer paper to 2026.
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text for cell in row.cells)
            if "A Phase-Enhanced Neural Network With Dual-Path Transformer" not in row_text:
                continue
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "2025" in paragraph.text:
                        set_text(paragraph, paragraph.text.replace("2025", "2026"))


def apply_final_content_corrections(doc: Document) -> None:
    try:
        set_text(find_paragraph(doc, "Date: 07: 07: 2026"), "Date: 07/07/2026")
    except ValueError:
        pass

    problem_two = find_startswith(doc, "Heart and lung sounds are also not easy")
    replace_inline_text(
        problem_two,
        "Heart and lung sounds are also not easy to differentiate since",
        "Heart and lung sounds are difficult to differentiate because",
    )
    replace_inline_text(
        problem_two,
        "a set frequency filter could filters out information or introduce sound that is not required",
        "a fixed frequency filter could remove useful information or retain unwanted sound",
    )
    replace_inline_text(
        problem_two,
        "In the case of a single mixed recording, without separate original heart and lung sources available, the problem is more difficult to solve than the prototype.",
        "With only one mixed recording and no isolated source signals, separation becomes more difficult because the system has limited information about the original heart and lung components.",
    )

    problem_three = find_startswith(doc, "A third difficulty is that there is no clear software")
    replace_inline_text(problem_three, "software re-usable process", "reusable software workflow")
    replace_inline_text(
        problem_three,
        "The majority of related studies are diagnosis, classification or algorithm level experiment.",
        "Many related studies focus on diagnosis, classification, or algorithm-level experiments.",
    )
    replace_inline_text(
        problem_three,
        "The prototype must be based on the following requirements: clear requirement, responsibilities of modules, flow of the interface, planning of storage and planning of evaluation, so as to be able to be implemented in FYP2 in a controlled manner.",
        "The prototype therefore needs clear requirements, defined module responsibilities, an interface workflow, a storage plan and an evaluation plan so that it can be implemented systematically during FYP2.",
    )

    for paragraph in doc.paragraphs:
        for old, new in (
            ("super relevant", "directly relevant"),
            ("Ullah and Zhang, 2024", "Ullah & Zhang, 2024"),
        ):
            if old in paragraph.text:
                replace_inline_text(paragraph, old, new)

    set_text(
        find_startswith(doc, "The questionnaire consists of 24 items"),
        "The full questionnaire contains 24 items and is provided in Appendix H. Chapter 3 analyses 10 requirement-focused items to avoid repeating closely related questions while preserving traceability between the questionnaire and the resulting requirements.",
    )
    set_text(
        find_startswith(doc, "Section A requested for the respondent"),
        "Section A collected respondent background information, including age group, respondent category, familiarity with audio or signal processing, biomedical-audio experience and experience with audio-related tasks. These responses describe the sample, so they are summarised in text rather than presented as separate figures.",
    )
    set_text(
        find_startswith(doc, "With respect to familiarity with audio or signal processing"),
        "With respect to familiarity with audio or signal processing, 29 respondents (54.7%) selected basic, 15 (28.3%) selected not familiar, 6 (11.3%) selected intermediate and 3 (5.7%) selected advanced. Thirty respondents (56.6%) reported prior biomedical-audio experience, 22 (41.5%) reported no prior experience and 1 (1.9%) was unsure. The sample therefore provides useful feedback from technical and student users, including some health-related respondents, but it cannot be treated as representative of hospital personnel. That limitation affects how confidently the requirements can be generalised to a clinic workflow.",
    )
    set_text(
        find_startswith(doc, "Finding: From the 53 respondents that answered this question"),
        "Finding: Responses were mixed. The largest group, 17 respondents (32.1%), selected neutral. A total of 21 respondents (39.6%) selected easy or very easy, while 15 (28.3%) selected difficult or very difficult. These results do not show a clear agreement that separation is either simple or difficult from the user's perspective. They do, however, support keeping the technical separation process behind a straightforward application workflow.",
    )
    set_text(
        find_startswith(doc, "Requirement implication: This finding implies that the system design should hide"),
        "Requirement implication: The interface should provide a streamlined process that allows users to upload a mixed cardiopulmonary recording and receive separate heart and lung outputs without requiring them to configure the internal signal-processing steps.",
    )
    set_text(
        find_startswith(doc, "Finding: 22 or 41.5%"),
        "Finding: Twenty-two respondents (41.5%) selected agree and 8 (15.1%) selected strongly agree. Together, 30 respondents (56.6%) gave a positive response. The remaining 23 respondents (43.4%) selected neutral, disagree or strongly disagree. The results therefore show more positive than non-positive responses, while still leaving a substantial group that was neutral or unconvinced.",
    )
    set_text(
        find_startswith(doc, "Finding: Most respondents (19 or 35.8%)"),
        "Finding: The largest group, 19 respondents (35.8%), rated audio preview as important. Another 12 (22.6%) selected very important and 12 (22.6%) selected moderately important. Eight respondents (15.1%) selected slightly important, while 2 (3.8%) selected not important.",
    )
    set_text(
        find_startswith(doc, "Finding: The majority of the respondents (22 or 41.5%)"),
        "Finding: The largest group, 22 respondents (41.5%), preferred a detailed interface with many options. A balanced interface with a simple view and optional advanced settings was selected by 20 respondents (37.7%), while 11 (20.8%) preferred a simple interface with only basic options. In a related question, 33 respondents (62.3%) rated ease of use as moderately important, important or very important. For error feedback, 29 respondents (54.7%) preferred a message that explains the problem and suggests a correction.",
    )
    set_text(
        find_startswith(doc, "Requirement implication: The cardiopulmonary separation prototype should focus on providing a good balance"),
        "Requirement implication: The prototype should provide a clear default workflow with advanced controls available when needed. Model selection, preprocessing information, evaluation results, waveforms and spectrograms can be shown as additional controls without crowding the main upload and result workflow.",
    )
    set_text(
        find_startswith(doc, "Finding: 16 or 30.2%"),
        "Finding: Sixteen respondents (30.2%) rated evaluation metrics as important, 11 (20.8%) selected moderately important and 8 (15.1%) selected very important. Eight respondents (15.1%) selected slightly important and 10 (18.9%) selected not important. A related question asked which information would be useful and listed signal quality, separation quality, noise reduction, processing time and an all-of-the-above option.",
    )

    set_text(
        find_startswith(doc, "The front end is supposed to be intuitive"),
        "The interface should make the separation-only scope clear and provide understandable status information. It should also display useful messages for unsupported files, unavailable models, preprocessing failures and separation failures.",
    )
    set_text(
        find_startswith(doc, "The database design follows the same principles as the Software Design project"),
        "The database design uses relational tables for metadata and file-system storage for binary audio. SQLite is suitable for the first local prototype because it is lightweight and easy to inspect during iterative development. Audio records, model definitions, processing jobs and separation results remain linked so that uploads and outputs can be traced.",
    )
    set_text(
        find_startswith(doc, "The prototype will utilize light-weight database storage"),
        "The prototype will use a lightweight database for metadata and file-system storage for audio. The database should store upload attributes, model metadata, selected model ID, job status, output paths, processing time, optional evaluation metrics and system logs. The file system will store uploaded recordings, model files and generated outputs.",
    )

    set_text(
        first_text_paragraph_after(doc, "6.1 Overview"),
        "This chapter concludes the FYP1 report by summarising the background study, literature review, requirements analysis, system design and implementation plan prepared for the cardiopulmonary sound separation prototype.",
    )
    set_text(
        find_startswith(doc, "The cardiopulmonary sound separation prototype to be developed during FYP1"),
        "The planned FYP2 workflow covers audio upload, validation, preprocessing, model selection, separation, preview, download and processing history. The FYP1 requirements and design define how these functions should work together before implementation begins.",
    )
    set_text(
        find_startswith(doc, "Regarding the expected contribution within the scope of FYP1"),
        "At the FYP1 stage, the completed contribution consists of the background study, requirements analysis, design artefacts and implementation plan needed to guide development and evaluation during FYP2.",
    )
    set_text(
        find_startswith(doc, "I wanted to make it clear that the intended contribution"),
        "The prototype is limited to separating mixed cardiopulmonary recordings into heart and lung sound outputs. It will not provide a diagnosis, treatment recommendation or clinical decision.",
    )
    set_text(
        first_text_paragraph_after(doc, "6.6 Summary"),
        "The FYP1 work defines an application-based prototype with two user roles, a traceable audio-record workflow, selectable separation strategies and a practical FYP2 implementation plan. Implementation, target-user validation, testing and performance evaluation remain future work and are not claimed as completed results.",
    )


def parse_bib_entries(text: str) -> dict[str, str]:
    entries: dict[str, str] = {}
    for part in re.split(r"(?m)(?=^@)", text):
        match = re.match(r"@\w+\{([^,]+),", part.strip())
        if match:
            entries[match.group(1)] = part.strip() + "\n"
    return entries


def source_reference_keys(doc: Document) -> list[str]:
    refs_heading = find_paragraph(doc, "References", last=True)
    appendix = find_paragraph(doc, "Appendix A: Gantt Chart", last=True)
    appendix.paragraph_format.page_break_before = True
    blocks = list(iter_block_items(doc))
    start = next(i for i, b in enumerate(blocks) if isinstance(b, Paragraph) and b._element is refs_heading._element)
    end = next(i for i, b in enumerate(blocks) if i > start and isinstance(b, Paragraph) and b._element is appendix._element)
    source_refs = [b.text.strip() for b in blocks[start + 1 : end] if isinstance(b, Paragraph) and b.text.strip()]

    bib_path = ROOT / "literature-review" / "references" / "references.bib"
    bib_text = bib_path.read_text(encoding="utf-8")
    entries = parse_bib_entries(bib_text)
    doi_to_key: dict[str, str] = {}
    for key, entry in entries.items():
        doi = re.search(r"doi\s*=\s*[\{\"]([^\}\"]+)", entry, re.I)
        if doi:
            doi_to_key[doi.group(1).lower().rstrip(".")] = key

    keys: list[str] = []
    for reference in source_refs:
        doi = re.search(r"https?://doi\.org/([^\s]+)", reference, re.I)
        if not doi:
            raise ValueError(f"Reference has no DOI and cannot be matched safely: {reference}")
        key = doi_to_key.get(doi.group(1).lower().rstrip("."))
        if not key:
            raise ValueError(f"Reference DOI not found in BibTeX: {doi.group(1)}")
        keys.append(key)
    return keys


def generate_apa_reference_docx(keys: list[str]) -> Path:
    APA_WORK.mkdir(parents=True, exist_ok=True)
    main_bib = ROOT / "report" / "revisions" / "verified_references.bib"
    entries = parse_bib_entries(main_bib.read_text(encoding="utf-8"))
    selected = [entries[key] for key in keys]
    selected.extend(
        [
            """@misc{threeMCore,
  author = {{3M}},
  title = {{3M} {Littmann} {CORE} digital stethoscope},
  url = {https://www.littmann.com/en-us/home/f/b5005222000/},
  urldate = {2026-07-31}
}
""",
            """@misc{eko2025app,
  author = {{Eko Health}},
  title = {Overview of the {Eko} app},
  year = {2025},
  month = {7},
  day = {22},
  url = {https://support.ekohealth.com/hc/en-us/articles/10303967065883-Overview-of-the-Eko-App}
}
""",
            """@misc{thinklabsWave,
  author = {{Thinklabs Medical LLC}},
  title = {{Thinklabs Wave} app},
  url = {https://www.thinklabs.com/thinklabs-wave-app-1},
  urldate = {2026-07-31}
}
""",
        ]
    )
    filtered_bib = APA_WORK / "selected_references.bib"
    filtered_bib.write_text("\n".join(selected), encoding="utf-8")
    md = APA_WORK / "references.md"
    md.write_text(
        "---\n"
        f"bibliography: {filtered_bib.as_posix()}\n"
        f"csl: {(ROOT / 'report' / 'revisions' / 'apa-7th-edition.csl').as_posix()}\n"
        "nocite: '@*'\n"
        "---\n\n# References\n",
        encoding="utf-8",
    )
    output = APA_WORK / "apa_references.docx"
    command = ["quarto", "pandoc", str(md), "--citeproc", "-s", "-o", str(output)]
    result = subprocess.run(command, cwd=ROOT, text=True, capture_output=True)
    if result.returncode != 0 or not output.exists():
        raise RuntimeError(f"APA reference generation failed:\n{result.stdout}\n{result.stderr}")
    return output


def copy_reference_paragraph(anchor: Paragraph, source: Paragraph) -> Paragraph:
    target = anchor.insert_paragraph_before(style="Normal")
    target.paragraph_format.left_indent = Inches(0.5)
    target.paragraph_format.first_line_indent = Inches(-0.5)
    target.paragraph_format.line_spacing = 2
    target.paragraph_format.space_before = Pt(0)
    target.paragraph_format.space_after = Pt(0)
    target.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for item in source.iter_inner_content():
        runs = item.runs if hasattr(item, "runs") else [item]
        for source_run in runs:
            run = target.add_run(source_run.text)
            run.bold = source_run.bold
            run.italic = source_run.italic
            run.underline = source_run.underline
            run.font.name = "Arial"
            run.font.size = Pt(11)
    return target


def update_references_to_apa(doc: Document) -> None:
    keys = source_reference_keys(doc)
    apa_docx = generate_apa_reference_docx(keys)
    apa_doc = Document(apa_docx)
    heading_index = next(i for i, p in enumerate(apa_doc.paragraphs) if p.text.strip() == "References")
    references = [p for p in apa_doc.paragraphs[heading_index + 1 :] if p.text.strip()]
    expected = len(keys) + 3
    if len(references) != expected:
        raise ValueError(f"Expected {expected} APA references, generated {len(references)}")

    refs_heading = find_paragraph(doc, "References", last=True)
    appendix = find_paragraph(doc, "Appendix A: Gantt Chart", last=True)
    blocks = list(iter_block_items(doc))
    start = next(i for i, b in enumerate(blocks) if isinstance(b, Paragraph) and b._element is refs_heading._element)
    end = next(i for i, b in enumerate(blocks) if i > start and isinstance(b, Paragraph) and b._element is appendix._element)
    reference_section = None
    for block in blocks[start + 1 : end]:
        if not isinstance(block, Paragraph) or block._p.pPr is None:
            continue
        section = block._p.pPr.find(qn("w:sectPr"))
        if section is not None:
            reference_section = deepcopy(section)
    if reference_section is None:
        raise ValueError("Reference section break was not found in the source report")

    for block in blocks[start + 1 : end]:
        block._element.getparent().remove(block._element)
    inserted_references: list[Paragraph] = []
    for reference in references:
        inserted_references.append(copy_reference_paragraph(appendix, reference))

    # Reattach the source section break that ended the References section. Without
    # this break, Word applies the following landscape appendix layout to References.
    page_number = reference_section.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        reference_section.append(page_number)
    page_number.set(qn("w:fmt"), "decimal")
    page_number.attrib.pop(qn("w:start"), None)
    ppr = inserted_references[-1]._p.get_or_add_pPr()
    existing_section = ppr.find(qn("w:sectPr"))
    if existing_section is not None:
        ppr.remove(existing_section)
    ppr.append(reference_section)


def apply_main_page_numbering(doc: Document) -> None:
    sections = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            ppr = child.find(qn("w:pPr"))
            section = ppr.find(qn("w:sectPr")) if ppr is not None else None
            if section is not None:
                sections.append(section)
    final_section = doc.element.body.find(qn("w:sectPr"))
    if final_section is not None:
        sections.append(final_section)

    main_index = None
    for index, section in enumerate(sections):
        page_number = section.find(qn("w:pgNumType"))
        if page_number is not None and page_number.get(qn("w:start")) == "1":
            main_index = index
            break
    if main_index is None:
        raise ValueError("The Chapter 1 page-number section was not found")

    for index, section in enumerate(sections[main_index:], start=main_index):
        page_number = section.find(qn("w:pgNumType"))
        if page_number is None:
            page_number = OxmlElement("w:pgNumType")
            section.append(page_number)
        page_number.set(qn("w:fmt"), "decimal")
        if index == main_index:
            page_number.set(qn("w:start"), "1")
        else:
            page_number.attrib.pop(qn("w:start"), None)


def add_seq_field(paragraph: Paragraph, kind: str, result: int, reset: bool) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    switch = f" \\r {result}" if reset else ""
    instr.text = f" SEQ {kind} \\* ARABIC{switch} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = str(result)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def apply_chapter_caption_numbering(doc: Document) -> None:
    chapter: int | None = None
    counters: dict[tuple[int, str], int] = {}
    label_map: dict[str, str] = {}
    caption_data: list[tuple[Paragraph, str, int, int, str, str]] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().replace("\xa0", " ")
        chapter_match = re.match(r"CHAPTER\s+([1-6]):", text)
        if paragraph.style.name == "Heading 1" and chapter_match:
            chapter = int(chapter_match.group(1))
        if paragraph.style.name != "Caption" or chapter is None:
            continue
        match = re.match(r"^(Table|Figure)\s+([0-9]+(?:\.[0-9]+)?):\s*(.+)$", text)
        if not match:
            continue
        kind, old_no, caption_text = match.groups()
        key = (chapter, kind)
        counters[key] = counters.get(key, 0) + 1
        number = counters[key]
        old_label = f"{kind} {old_no}"
        new_label = f"{kind} {chapter}.{number}"
        label_map[old_label] = new_label
        caption_data.append((paragraph, kind, chapter, number, caption_text, new_label))

    if label_map:
        labels = sorted(label_map, key=len, reverse=True)
        pattern = re.compile(r"\b(" + "|".join(re.escape(label) for label in labels) + r")(?![.0-9])")
        for paragraph in doc.paragraphs:
            style = paragraph.style.name.lower()
            if paragraph.style.name == "Caption" or style.startswith("toc") or style == "table of figures":
                continue
            old = paragraph.text
            new = pattern.sub(lambda match: label_map[match.group(1)], old)
            if new != old:
                set_text(paragraph, new)

    for paragraph, kind, chapter_no, number, caption_text, _ in caption_data:
        paragraph.clear()
        paragraph.add_run(f"{kind} {chapter_no}.")
        add_seq_field(paragraph, kind, number, reset=(number == 1))
        paragraph.add_run(f": {caption_text}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True


def apply_justified_alignment(doc: Document) -> None:
    in_body = False
    in_front_matter_prose = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Acknowledgements":
            in_front_matter_prose = True
        if text == "Table of Contents":
            in_front_matter_prose = False
        if text == "CHAPTER 1: INTRODUCTION":
            in_body = True
        if text == "References" and paragraph.style.name == "Heading 1":
            in_body = False
        if text == "Appendix A: Gantt Chart" and paragraph.style.name == "Heading 1":
            in_body = False
        if paragraph.style.name == "Heading 1" and text.startswith("CHAPTER "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if not (in_body or in_front_matter_prose) or not text or paragraph.style.name != "Normal":
            continue
        if paragraph._element.xpath(".//m:oMath | .//m:oMathPara | .//w:drawing | .//w:pict"):
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def request_field_update_on_open(doc: Document) -> None:
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def png_dimensions(data: bytes) -> tuple[int, int]:
    if data[:8] != b"\x89PNG\r\n\x1a\n":
        raise ValueError("Diagram is not a PNG file")
    return struct.unpack(">II", data[16:24])


def replace_diagrams_and_fit(docx_path: Path) -> None:
    replacements = {
        "word/media/image14.png": ROOT / "report" / "quarto" / "figures" / "plantuml" / "context_diagram.png",
        "word/media/image15.png": ROOT / "report" / "quarto" / "figures" / "plantuml" / "use_case_diagram.png",
        "word/media/image16.png": ROOT / "report" / "quarto" / "figures" / "plantuml" / "activity_diagram.png",
        "word/media/image17.png": ROOT / "report" / "quarto" / "figures" / "plantuml" / "class_diagram.png",
        "word/media/image18.png": ROOT / "report" / "quarto" / "figures" / "plantuml" / "component_diagram.png",
        "word/media/image19.png": ROOT / "report" / "quarto" / "figures" / "plantuml" / "sequence_diagram.png",
    }
    from lxml import etree

    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    with zipfile.ZipFile(docx_path, "r") as zin:
        content = {item.filename: zin.read(item.filename) for item in zin.infolist()}

    rels = etree.fromstring(content["word/_rels/document.xml.rels"])
    rid_to_target = {
        rel.get("Id"): "word/" + rel.get("Target").replace("\\", "/")
        for rel in rels.findall("pr:Relationship", ns)
        if rel.get("Target", "").startswith("media/")
    }
    document = etree.fromstring(content["word/document.xml"])
    max_width = int(6.35 * 914400)
    max_height = int(7.8 * 914400)

    for blip in document.xpath(".//a:blip", namespaces=ns):
        rid = blip.get(f"{{{ns['r']}}}embed")
        target = rid_to_target.get(rid)
        if target not in replacements:
            continue
        image_data = replacements[target].read_bytes()
        width_px, height_px = png_dimensions(image_data)
        drawing = blip
        while drawing is not None and drawing.tag != f"{{{ns['w']}}}drawing":
            drawing = drawing.getparent()
        if drawing is None:
            continue
        wp_extent = drawing.find(".//wp:extent", ns)
        a_extent = drawing.find(".//a:xfrm/a:ext", ns)
        original_width = int(wp_extent.get("cx")) if wp_extent is not None else max_width
        width = min(original_width, max_width)
        height = int(width * height_px / width_px)
        if height > max_height:
            height = max_height
            width = int(height * width_px / height_px)
        if wp_extent is not None:
            wp_extent.set("cx", str(width))
            wp_extent.set("cy", str(height))
        if a_extent is not None:
            a_extent.set("cx", str(width))
            a_extent.set("cy", str(height))
        content[target] = image_data

    content["word/document.xml"] = etree.tostring(
        document, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in content.items():
            zout.writestr(name, data)
    tmp.replace(docx_path)


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUT)
    doc = Document(OUT)

    replace_abstract(doc)
    update_application_framing(doc)
    add_chapter2_application_study(doc)
    update_requirements_chapter(doc)
    update_design_and_plan(doc)
    apply_final_content_corrections(doc)
    update_references_to_apa(doc)
    apply_main_page_numbering(doc)
    apply_chapter_caption_numbering(doc)
    apply_justified_alignment(doc)
    request_field_update_on_open(doc)
    doc.save(OUT)
    replace_diagrams_and_fit(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
